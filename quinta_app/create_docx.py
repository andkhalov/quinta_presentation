#!/usr/bin/env python3
"""
Create DOCX presentation document for Quinta Essentia Marketing Plan.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = '../presentation_materials'
os.makedirs(OUTPUT_DIR, exist_ok=True)

def create_presentation_docx():
    """Create the main presentation document."""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title style
    title_style = doc.styles.add_style('SlideTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Arial'
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(184, 134, 11)  # Gold
    
    # Subtitle style  
    subtitle_style = doc.styles.add_style('SlideSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_style.font.name = 'Arial'
    subtitle_style.font.size = Pt(14)
    subtitle_style.font.color.rgb = RGBColor(100, 100, 100)
    
    def add_slide(title, content_func):
        """Add a slide-like section."""
        doc.add_paragraph('─' * 50)
        p = doc.add_paragraph(title, style='SlideTitle')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        content_func()
        doc.add_page_break()
    
    # ============================================
    # SLIDE 1: Title
    # ============================================
    def slide1_content():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('\n\n\n🌟 QUINTA ESSENTIA 🌟\n\n')
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(184, 134, 11)
        
        p2 = doc.add_paragraph('Партнёрская программа')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(18)
        
        p3 = doc.add_paragraph('\n\nПуть к финансовой свободе\nчерез здоровье и благополучие')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.runs[0].font.size = Pt(14)
        p3.runs[0].font.color.rgb = RGBColor(100, 100, 100)
        
        p4 = doc.add_paragraph('\n\nquinta.pro')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.runs[0].font.size = Pt(12)
        p4.runs[0].font.italic = True
    
    add_slide('', slide1_content)
    
    # ============================================
    # SLIDE 2: Why we're here
    # ============================================
    def slide2_content():
        doc.add_paragraph('Три причины обратить внимание:', style='SlideSubtitle')
        doc.add_paragraph()
        
        reasons = [
            ('1. Продукт, который работает', 'Научно обоснованные формулы для здоровья'),
            ('2. Честный заработок', '5 источников дохода с первого дня'),
            ('3. Система поддержки', 'Вы никогда не одиноки на этом пути'),
        ]
        
        for title, desc in reasons:
            p = doc.add_paragraph()
            run = p.add_run(f'✅ {title}')
            run.font.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(f'    {desc}')
            doc.add_paragraph()
        
        quote = doc.add_paragraph()
        quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = quote.add_run('\n«Я не продаю — я делюсь тем, что изменило мою жизнь»')
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    add_slide('ПОЧЕМУ МЫ ЗДЕСЬ', slide2_content)
    
    # ============================================
    # SLIDE 3: Trust
    # ============================================
    def slide3_content():
        doc.add_paragraph('Quinta Essentia — это надёжно:', style='SlideSubtitle')
        doc.add_paragraph()
        
        points = [
            'Европейская компания с прозрачной структурой',
            'Продукция прошла клинические исследования',
            'Более 50 000 довольных клиентов',
            'Работаем в 15+ странах мира',
            'Полная юридическая чистота',
        ]
        
        for point in points:
            p = doc.add_paragraph(f'✅ {point}')
            p.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
        p = doc.add_paragraph('Мы не обещаем чудес — мы даём инструменты')
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_slide('ДОВЕРИЕ К КОМПАНИИ', slide3_content)
    
    # ============================================
    # SLIDE 4: How much can you earn
    # ============================================
    def slide4_content():
        doc.add_paragraph('Реальные цифры дохода:', style='SlideSubtitle')
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        
        headers = ['Стратегия', '6 месяцев', '12 месяцев']
        data = [
            ['🛒 Клиент', '€50-100', '€100-200'],
            ['⚖️ Партнёр', '€500-1000', '€1000-2000'],
            ['🏗️ Строитель', '€1500-3000', '€3000-6000'],
            ['👑 Лидер', '€3000-5000', '€10000+'],
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = cell_text
        
        doc.add_paragraph()
        p = doc.add_paragraph('Выбор за вами — от небольшого бонуса до полной финансовой свободы')
        p.runs[0].font.italic = True
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: income_growth.png]')
    
    add_slide('СКОЛЬКО МОЖНО ЗАРАБОТАТЬ', slide4_content)
    
    # ============================================
    # SLIDE 5: 5 Income Sources
    # ============================================
    def slide5_content():
        doc.add_paragraph('Вы зарабатываете сразу из 5 источников:', style='SlideSubtitle')
        doc.add_paragraph()
        
        sources = [
            ('🎁 КЕШБЭК', '5-12.5%', 'Возврат с личных покупок'),
            ('💼 БОНУС С ПРОДАЖ', '2.5-7.5%', 'С покупок приглашённых'),
            ('👥 КОМАНДНЫЙ', '1-5%', 'С оборота всей команды'),
            ('👑 ЛИДЕРСКИЙ', '2.5%×5', 'Дополнительно для лидеров'),
            ('🏆 TOP БОНУС', 'до 3%', 'Доля от оборота компании'),
        ]
        
        for name, rate, desc in sources:
            p = doc.add_paragraph()
            run = p.add_run(f'{name} ({rate})')
            run.font.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(f'    {desc}')
        
        doc.add_paragraph()
        p = doc.add_paragraph('✅ Все бонусы работают одновременно!')
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(45, 80, 22)
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: 5_income_sources.png]')
    
    add_slide('5 ИСТОЧНИКОВ ДОХОДА', slide5_content)
    
    # ============================================
    # SLIDE 6: Cashback
    # ============================================
    def slide6_content():
        doc.add_paragraph('Начните экономить с первого дня:', style='SlideSubtitle')
        doc.add_paragraph()
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ['Ваши покупки', 'Кешбэк'],
            ['35 PV (~€65)', '5%'],
            ['70 PV (~€130)', '7.5%'],
            ['140 PV (~€260)', '10%'],
            ['280 PV (~€520)', '12.5%'],
        ]
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        p = doc.add_paragraph('Пример: При покупках на 140 PV в месяц вы получаете 14 PV (~€8) просто так!')
        p.runs[0].font.italic = True
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: cashback_rates.png]')
    
    add_slide('КЕШБЭК — ПЕРВЫЕ ДЕНЬГИ', slide6_content)
    
    # ============================================
    # SLIDE 7: Team Bonus
    # ============================================
    def slide7_content():
        doc.add_paragraph('Глубина — ваш актив:', style='SlideSubtitle')
        doc.add_paragraph()
        
        levels = [
            ('1 уровень', '5%', 'прямые партнёры'),
            ('2-4 уровень', '2.5%', ''),
            ('5-7 уровень', '1.5%', ''),
            ('8+ уровень', '1%', 'бесконечно для Doctus+!'),
        ]
        
        for level, rate, note in levels:
            p = doc.add_paragraph()
            run = p.add_run(f'• {level}: ')
            run.font.bold = True
            run2 = p.add_run(rate)
            run2.font.color.rgb = RGBColor(184, 134, 11)
            run2.font.bold = True
            if note:
                p.add_run(f' ({note})')
        
        doc.add_paragraph()
        p = doc.add_paragraph('Чем глубже структура — тем больше пассивный доход')
        p.runs[0].font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: team_bonus_depth.png]')
    
    add_slide('КОМАНДНЫЙ БОНУС', slide7_content)
    
    # ============================================
    # SLIDE 8: Ranks System
    # ============================================
    def slide8_content():
        doc.add_paragraph('12 ступеней к вершине:', style='SlideSubtitle')
        doc.add_paragraph()
        
        ranks = [
            '🌱 Novus', '🌿 Inceptor', '🌲 Cognitor', '⭐ Doctus',
            '✨ Primum', '💎 Dux', '👑 Provectus', '🏆 Grandis',
            '🌟 Oraculi', '💫 Sapiens', '🔮 Superator', '👸 Principal'
        ]
        
        p = doc.add_paragraph()
        for i, rank in enumerate(ranks):
            p.add_run(rank)
            if i < len(ranks) - 1:
                p.add_run(' → ')
        
        doc.add_paragraph()
        doc.add_paragraph('Каждый ранг открывает новые возможности:')
        doc.add_paragraph('• Глубже командный бонус')
        doc.add_paragraph('• Лидерский бонус')
        doc.add_paragraph('• Подарки и награды')
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: ranks_ladder.png]')
    
    add_slide('СИСТЕМА РАНГОВ', slide8_content)
    
    # ============================================
    # SLIDE 9: Rewards
    # ============================================
    def slide9_content():
        doc.add_paragraph('Компания щедро награждает лидеров:', style='SlideSubtitle')
        doc.add_paragraph()
        
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ['Ранг', 'Награда'],
            ['⭐ Doctus', '€500 (единоразово)'],
            ['✨ Primum', '€2000 (единоразово)'],
            ['💎 Dux', '€1000/месяц'],
            ['👑 Provectus', '€2000/месяц'],
            ['🏆 Grandis', 'АВТОМОБИЛЬ или €3500/мес'],
            ['💫 Sapiens', 'Премиум авто + TOP бонус'],
        ]
        
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: rewards.png]')
    
    add_slide('НАГРАДЫ ЗА РАНГИ', slide9_content)
    
    # ============================================
    # SLIDE 10: Strategies
    # ============================================
    def slide10_content():
        doc.add_paragraph('4 пути — один результат:', style='SlideSubtitle')
        doc.add_paragraph()
        
        strategies = [
            ('🛒 Клиент', '30-60 минут в день', '€100-200/мес'),
            ('⚖️ Партнёр', '2-3 часа в день', '€1000-2000/мес'),
            ('🏗️ Строитель', '4-6 часов в день', '€3000-6000/мес'),
            ('👑 Лидер', 'Полная занятость', '€10000+/мес'),
        ]
        
        for name, time, income in strategies:
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.font.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph(f'    ⏰ {time}')
            p2 = doc.add_paragraph(f'    💰 {income}')
            p2.runs[0].font.color.rgb = RGBColor(184, 134, 11)
            doc.add_paragraph()
        
        doc.add_paragraph('[Вставить график: strategies_comparison.png]')
    
    add_slide('КАКАЯ СТРАТЕГИЯ ПОДХОДИТ ВАМ?', slide10_content)
    
    # ============================================
    # SLIDE 11: For Everyone
    # ============================================
    def slide11_content():
        doc.add_paragraph('Кто уже добивается успеха:', style='SlideSubtitle')
        doc.add_paragraph()
        
        profiles = [
            ('👩‍💼 Мамы в декрете', 'работают из дома, пока дети спят'),
            ('👨‍🏫 Специалисты', 'дополнительный доход без смены работы'),
            ('👴 Пенсионеры', 'активная жизнь и прибавка к пенсии'),
            ('👨‍🎓 Студенты', 'гибкий график, первый бизнес-опыт'),
            ('👩‍⚕️ Врачи', 'монетизация рекомендаций'),
        ]
        
        for profile, desc in profiles:
            p = doc.add_paragraph()
            run = p.add_run(f'{profile} — ')
            run.font.bold = True
            p.add_run(desc)
        
        doc.add_paragraph()
        p = doc.add_paragraph('Общее у всех — желание изменить свою жизнь')
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(45, 80, 22)
    
    add_slide('ПОДХОДИТ ВСЕМ', slide11_content)
    
    # ============================================
    # SLIDE 12: First Steps
    # ============================================
    def slide12_content():
        doc.add_paragraph('От регистрации до первого дохода:', style='SlideSubtitle')
        doc.add_paragraph()
        
        steps = [
            ('1', '📝 Регистрация', 'Бесплатно, 5 минут'),
            ('2', '🛒 Первый заказ', '70 PV (≈€130)'),
            ('3', '💬 Рекомендации', 'Расскажите 2-3 друзьям'),
            ('4', '👥 Приглашения', 'Помогите 2 людям начать'),
            ('5', '💰 Первый доход', '€100-300 в первый месяц!'),
        ]
        
        for num, title, desc in steps:
            p = doc.add_paragraph()
            run = p.add_run(f'{num}. {title}')
            run.font.bold = True
            run.font.size = Pt(12)
            doc.add_paragraph(f'    {desc}')
        
        doc.add_paragraph()
        doc.add_paragraph('[Вставить график: first_steps.png]')
    
    add_slide('ПЕРВЫЕ ШАГИ', slide12_content)
    
    # ============================================
    # SLIDE 13: Training System
    # ============================================
    def slide13_content():
        doc.add_paragraph('Вы никогда не останетесь одни:', style='SlideSubtitle')
        doc.add_paragraph()
        
        support = [
            ('📚 База знаний', 'видеоуроки, инструкции, скрипты'),
            ('👥 Наставник', 'персональная поддержка 24/7'),
            ('🎓 Вебинары', 'обучение от топ-лидеров'),
            ('💬 Чаты команды', 'мгновенная помощь'),
            ('🏆 Конференции', 'мотивация и нетворкинг'),
        ]
        
        for icon_title, desc in support:
            p = doc.add_paragraph()
            run = p.add_run(f'{icon_title} — ')
            run.font.bold = True
            p.add_run(desc)
        
        doc.add_paragraph()
        p = doc.add_paragraph('Ваш успех — наш успех')
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_slide('СИСТЕМА ОБУЧЕНИЯ', slide13_content)
    
    # ============================================
    # SLIDE 14: Objections
    # ============================================
    def slide14_content():
        doc.add_paragraph('Ответы на частые вопросы:', style='SlideSubtitle')
        doc.add_paragraph()
        
        objections = [
            ('"У меня нет опыта в продажах"', 'Мы не продаём — мы рекомендуем. И этому легко научиться.'),
            ('"Это MLM?"', 'Да, но легальный. Вы зарабатываете на реальных продажах продукции, а не на взносах.'),
            ('"Нужны большие вложения?"', 'Минимум — €65/мес на продукт для себя. Никаких закупок.'),
            ('"Я интроверт"', 'Онлайн-инструменты позволяют работать без "холодных" контактов.'),
        ]
        
        for question, answer in objections:
            p = doc.add_paragraph()
            run = p.add_run(f'❓ {question}')
            run.font.italic = True
            p2 = doc.add_paragraph(f'→ {answer}')
            p2.runs[0].font.color.rgb = RGBColor(45, 80, 22)
            doc.add_paragraph()
    
    add_slide('ОТВЕТЫ НА ВОЗРАЖЕНИЯ', slide14_content)
    
    # ============================================
    # SLIDE 15: Call to Action
    # ============================================
    def slide15_content():
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run('\n\nДВА ПУТИ\n\n')
        run.font.size = Pt(18)
        run.font.bold = True
        
        p2 = doc.add_paragraph('Путь 1: Закрыть эту презентацию и забыть.')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Через год ничего не изменится.').alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        p3 = doc.add_paragraph('Путь 2: Попробовать.')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.runs[0].font.bold = True
        p4 = doc.add_paragraph('Через год вы можете быть финансово свободны.')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.runs[0].font.color.rgb = RGBColor(184, 134, 11)
        
        doc.add_paragraph('\n\n')
        
        p5 = doc.add_paragraph('НАЧНИТЕ ПРЯМО СЕЙЧАС')
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.runs[0].font.size = Pt(16)
        p5.runs[0].font.bold = True
        
        p6 = doc.add_paragraph('quinta.pro')
        p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p6.runs[0].font.size = Pt(14)
        
        doc.add_paragraph('\n')
        
        quote = doc.add_paragraph('«Лучшее время посадить дерево было 20 лет назад.\nВторое лучшее время — сейчас»')
        quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
        quote.runs[0].font.italic = True
        quote.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    
    add_slide('ФИНАЛЬНЫЙ ПРИЗЫВ', slide15_content)
    
    # Save document
    output_path = os.path.join(OUTPUT_DIR, 'Quinta_Essentia_Presentation.docx')
    doc.save(output_path)
    print(f'✓ Created {output_path}')
    return output_path


def main():
    print('\n' + '='*50)
    print('Creating DOCX presentation for Quinta Essentia')
    print('='*50 + '\n')
    
    create_presentation_docx()
    
    print('\n' + '='*50)
    print('Document created successfully!')
    print('='*50 + '\n')


if __name__ == '__main__':
    main()
